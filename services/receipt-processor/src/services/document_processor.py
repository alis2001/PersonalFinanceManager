"""
Comprehensive Document Processing Module - All Text-Based Files
Location: services/receipt-processor/src/services/document_processor.py

FOCUSED RESPONSIBILITY: Clean text extraction from ALL document types
- Plain text files (.txt, .rtf)
- Microsoft Office (.docx, .doc, .xlsx, .xls)
- CSV/TSV files (.csv, .tsv, .tab)
- JSON files (.json)
- XML files (.xml)
- Other structured data formats
- NO interpretation - just clean text extraction
"""

import time
import json
import xml.etree.ElementTree as ET
from typing import Dict, List, Optional, Any
from pathlib import Path
import zipfile
import tempfile

import pandas as pd
import openpyxl
import csv
import docx
from docx import Document
import structlog

from ..config.settings import settings

logger = structlog.get_logger(__name__)

class DocumentProcessor:
    """Universal text extractor for ALL document file types"""
    
    def __init__(self):
        self.supported_formats = {
            # Plain text
            '.txt', '.text', '.rtf',
            # Microsoft Office
            '.docx', '.doc', '.xlsx', '.xls',
            # Data formats
            '.csv', '.tsv', '.tab',
            # Structured data
            '.json', '.xml',
            # Other text formats
            '.log', '.md', '.markdown'
        }
        self.max_rows = 2000
        self.max_file_size = settings.MAX_FILE_SIZE_BYTES
    
    async def extract_text_from_document(self, job_id: str, file_path: str, filename: str) -> Dict[str, any]:
        """
        Extract RAW TEXT from any document file type
        
        Args:
            job_id: Processing job ID
            file_path: Path to document file
            filename: Original filename
            
        Returns:
            Dict with RAW extracted text only
        """
        start_time = time.time()
        
        try:
            file_ext = Path(filename).suffix.lower()
            
            logger.info("📄 Starting text extraction from document", 
                       job_id=job_id,
                       filename=filename,
                       file_type=file_ext)
            
            # Basic validation
            if not self._is_file_valid(file_path, file_ext):
                return {
                    "success": False,
                    "error": "Document file is invalid, too large, or unsupported"
                }
            
            # Route to appropriate extractor based on file type
            if file_ext in {'.txt', '.text', '.log', '.md', '.markdown'}:
                result = await self._extract_from_plaintext(file_path)
                
            elif file_ext == '.rtf':
                result = await self._extract_from_rtf(file_path)
                
            elif file_ext in {'.docx', '.doc'}:
                result = await self._extract_from_word(file_path, file_ext)
                
            elif file_ext in {'.xlsx', '.xls'}:
                result = await self._extract_from_excel(file_path, file_ext)
                
            elif file_ext in {'.csv', '.tsv', '.tab'}:
                result = await self._extract_from_csv(file_path, file_ext)
                
            elif file_ext == '.json':
                result = await self._extract_from_json(file_path)
                
            elif file_ext == '.xml':
                result = await self._extract_from_xml(file_path)
                
            else:
                # Try as plain text fallback
                logger.info("Unknown file type, attempting plain text extraction", file_ext=file_ext)
                result = await self._extract_from_plaintext(file_path)
            
            processing_time_ms = int((time.time() - start_time) * 1000)
            
            if result["success"]:
                logger.info("✅ Document text extraction completed", 
                           job_id=job_id,
                           method=result["method"],
                           text_length=len(result["text"]),
                           processing_time_ms=processing_time_ms)
            
            result["processing_time_ms"] = processing_time_ms
            return result
            
        except Exception as e:
            processing_time_ms = int((time.time() - start_time) * 1000)
            error_msg = f"Document extraction failed: {str(e)}"
            
            logger.error("❌ Document extraction failed", 
                        job_id=job_id,
                        error=str(e),
                        processing_time_ms=processing_time_ms)
            
            return {
                "success": False,
                "error": error_msg,
                "processing_time_ms": processing_time_ms
            }

    def _is_file_valid(self, file_path: str, file_ext: str) -> bool:
        """Basic file validation"""
        try:
            if not Path(file_path).exists():
                return False
            
            if file_ext not in self.supported_formats:
                return False
            
            file_size = Path(file_path).stat().st_size
            return file_size <= self.max_file_size
            
        except:
            return False

    async def _extract_from_plaintext(self, file_path: str) -> Dict[str, any]:
        """Extract from plain text files (.txt, .log, .md, etc.)"""
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # If all encodings fail, try with error handling
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
            
            # Basic cleanup - remove excessive whitespace only
            cleaned_text = self._basic_text_cleanup(text)
            
            return {
                "success": True,
                "text": cleaned_text,
                "confidence": 1.0,  # Perfect confidence for plain text
                "method": "plaintext"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Plain text extraction failed: {str(e)}"
            }

    async def _extract_from_rtf(self, file_path: str) -> Dict[str, any]:
        """Extract from RTF files"""
        try:
            # Simple RTF parsing - strip RTF codes and extract text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                rtf_content = file.read()
            
            # Basic RTF code removal
            text = self._strip_rtf_codes(rtf_content)
            cleaned_text = self._basic_text_cleanup(text)
            
            return {
                "success": True,
                "text": cleaned_text,
                "confidence": 0.9,
                "method": "rtf_basic"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"RTF extraction failed: {str(e)}"
            }

    async def _extract_from_word(self, file_path: str, file_ext: str) -> Dict[str, any]:
        """Extract from Word documents (.docx, .doc)"""
        try:
            if file_ext == '.docx':
                # Use python-docx for .docx files
                doc = Document(file_path)
                
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # Extract tables
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        text_parts.append(f"\n--- Table ---\n{table_text}")
                
                extracted_text = '\n'.join(text_parts)
                method = "python-docx"
                
            else:
                # For .doc files, try reading as plain text (fallback)
                logger.warning(".doc format not fully supported, attempting plain text extraction")
                return await self._extract_from_plaintext(file_path)
            
            cleaned_text = self._basic_text_cleanup(extracted_text)
            
            return {
                "success": True,
                "text": cleaned_text,
                "confidence": 0.95,
                "method": method
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Word document extraction failed: {str(e)}"
            }

    async def _extract_from_excel(self, file_path: str, file_ext: str) -> Dict[str, any]:
        """Extract from Excel files (.xlsx, .xls)"""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None, nrows=self.max_rows)
            
            text_parts = []
            
            for sheet_name, df in excel_data.items():
                if df.empty:
                    continue
                
                # Convert DataFrame to simple text representation
                sheet_text = f"=== Sheet: {sheet_name} ===\n"
                
                # Add headers
                headers = ' | '.join(str(col) for col in df.columns)
                sheet_text += f"{headers}\n"
                
                # Add data rows
                for _, row in df.iterrows():
                    row_values = []
                    for val in row.values:
                        if pd.notna(val):
                            row_values.append(str(val))
                        else:
                            row_values.append("")
                    
                    if any(val.strip() for val in row_values):
                        sheet_text += ' | '.join(row_values) + '\n'
                
                text_parts.append(sheet_text)
            
            extracted_text = '\n\n'.join(text_parts)
            cleaned_text = self._basic_text_cleanup(extracted_text)
            
            return {
                "success": True,
                "text": cleaned_text,
                "confidence": 0.98,
                "method": "pandas_excel"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Excel extraction failed: {str(e)}"
            }

    async def _extract_from_csv(self, file_path: str, file_ext: str) -> Dict[str, any]:
        """Extract from CSV/TSV files"""
        try:
            # Determine delimiter
            delimiter = {
                '.csv': ',',
                '.tsv': '\t',
                '.tab': '\t'
            }.get(file_ext, ',')
            
            # Read CSV
            df = pd.read_csv(
                file_path, 
                delimiter=delimiter,
                nrows=self.max_rows,
                encoding='utf-8',
                on_bad_lines='skip'
            )
            
            if df.empty:
                return {
                    "success": False,
                    "error": "CSV file contains no readable data"
                }
            
            # Convert to text
            text_parts = []
            
            # Add headers
            headers = ' | '.join(str(col) for col in df.columns)
            text_parts.append(headers)
            
            # Add data rows
            for _, row in df.iterrows():
                row_values = []
                for val in row.values:
                    if pd.notna(val):
                        row_values.append(str(val))
                    else:
                        row_values.append("")
                
                if any(val.strip() for val in row_values):
                    text_parts.append(' | '.join(row_values))
            
            extracted_text = '\n'.join(text_parts)
            cleaned_text = self._basic_text_cleanup(extracted_text)
            
            return {
                "success": True,
                "text": cleaned_text,
                "confidence": 0.95,
                "method": "pandas_csv"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"CSV extraction failed: {str(e)}"
            }

    async def _extract_from_json(self, file_path: str) -> Dict[str, any]:
        """Extract from JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text
            text = self._json_to_text(data)
            cleaned_text = self._basic_text_cleanup(text)
            
            return {
                "success": True,
                "text": cleaned_text,
                "confidence": 0.9,
                "method": "json_parser"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"JSON extraction failed: {str(e)}"
            }

    async def _extract_from_xml(self, file_path: str) -> Dict[str, any]:
        """Extract from XML files"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract all text content from XML
            text = self._xml_to_text(root)
            cleaned_text = self._basic_text_cleanup(text)
            
            return {
                "success": True,
                "text": cleaned_text,
                "confidence": 0.9,
                "method": "xml_parser"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"XML extraction failed: {str(e)}"
            }

    def _extract_table_text(self, table) -> str:
        """Extract text from Word table"""
        try:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_text.append(cell_text)
                
                if any(text for text in row_text):
                    table_text.append(' | '.join(row_text))
            
            return '\n'.join(table_text)
            
        except Exception as e:
            logger.warning("Table extraction failed", error=str(e))
            return ""

    def _strip_rtf_codes(self, rtf_content: str) -> str:
        """Basic RTF code stripping"""
        import re
        
        # Remove RTF control words
        text = re.sub(r'\\[a-z]+\d*\s?', '', rtf_content)
        
        # Remove RTF control symbols
        text = re.sub(r'\\[^a-z]', '', text)
        
        # Remove braces
        text = re.sub(r'[{}]', '', text)
        
        return text

    def _json_to_text(self, data: Any, level: int = 0) -> str:
        """Convert JSON data to readable text"""
        if level > 10:  # Prevent infinite recursion
            return "[max depth reached]"
        
        indent = "  " * level
        
        if isinstance(data, dict):
            text_parts = []
            for key, value in data.items():
                value_text = self._json_to_text(value, level + 1)
                text_parts.append(f"{indent}{key}: {value_text}")
            return '\n'.join(text_parts)
        
        elif isinstance(data, list):
            text_parts = []
            for i, item in enumerate(data):
                item_text = self._json_to_text(item, level + 1)
                text_parts.append(f"{indent}[{i}]: {item_text}")
            return '\n'.join(text_parts)
        
        else:
            return str(data)

    def _xml_to_text(self, element, level: int = 0) -> str:
        """Convert XML element to readable text"""
        if level > 20:  # Prevent infinite recursion
            return "[max depth reached]"
        
        text_parts = []
        indent = "  " * level
        
        # Add element tag and attributes
        if element.tag:
            tag_text = f"{indent}<{element.tag}>"
            if element.attrib:
                attrs = ' '.join(f'{k}="{v}"' for k, v in element.attrib.items())
                tag_text = f"{indent}<{element.tag} {attrs}>"
            text_parts.append(tag_text)
        
        # Add element text
        if element.text and element.text.strip():
            text_parts.append(f"{indent}  {element.text.strip()}")
        
        # Add children
        for child in element:
            child_text = self._xml_to_text(child, level + 1)
            text_parts.append(child_text)
        
        return '\n'.join(text_parts)

    def _basic_text_cleanup(self, text: str) -> str:
        """Basic text cleanup - remove excessive whitespace only"""
        if not text:
            return ""
        
        # Remove excessive whitespace but preserve structure
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Clean up line but preserve meaningful whitespace
            cleaned_line = ' '.join(line.split())
            cleaned_lines.append(cleaned_line)
        
        # Remove excessive empty lines
        final_text = '\n'.join(cleaned_lines)
        
        # Reduce multiple newlines to max 2
        import re
        final_text = re.sub(r'\n{3,}', '\n\n', final_text)
        
        return final_text.strip()

    def is_supported_format(self, file_extension: str) -> bool:
        """Check if file format is supported"""
        return file_extension.lower() in self.supported_formats

    def get_supported_formats(self) -> List[str]:
        """Get list of all supported file formats"""
        return sorted(list(self.supported_formats))